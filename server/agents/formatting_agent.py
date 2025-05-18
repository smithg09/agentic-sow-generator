import os
import docx
from docx.shared import Pt, Inches
import requests
from io import BytesIO
from config import LOGO_URL

class FormattingAgent:
    def __init__(self):
        pass
        
    def add_paragraph(self, doc, text, size=12):
        """Add a paragraph with specified font size (default 12pt)"""
        # Check if there are any markdown bold patterns (**text**)
        if "**" in text:
            paragraph = doc.add_paragraph()
            parts = text.split("**")
            
            # Parts will alternate between normal text and bold text
            for i, part in enumerate(parts):
                if part:  # Skip empty parts
                    # Even indices are normal text, odd indices are bold text
                    run = paragraph.add_run(part)
                    run.font.size = Pt(size)
                    if i % 2 == 1:  # This is text that was between ** and should be bold
                        run.bold = True
        else:
            # No markdown formatting, add as normal
            paragraph = doc.add_paragraph(text)
            for run in paragraph.runs:
                run.font.size = Pt(size)
        
        return paragraph
    
    def create_contractor_resources_table(self, doc, resources_data):
        # Create table with headers
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        
        # Set header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "Role"
        header_cells[1].text = "Resources"
        header_cells[2].text = "Responsibility"
        
        # Add data rows
        for resource in resources_data:
            row_cells = table.add_row().cells
            
            # Process role text for bold formatting
            role_text = resource.get('roleName', '')
            self.add_text_to_cell(row_cells[0], role_text)
            
            # Process resources text for bold formatting
            resources_text = str(resource.get('noOfPersons', ''))
            self.add_text_to_cell(row_cells[1], resources_text)
            
            # Process responsibility text for bold formatting
            responsibility_text = resource.get('responsibility', '')
            self.add_text_to_cell(row_cells[2], responsibility_text)
        
        return table
    
    def create_contractor_resources_markdown(self, resources_data):

        # Create markdown table header
        markdown_table = "| Role | Resources | Responsibility |\n"
        markdown_table += "|------|-----------------|---------------|\n"
        
        # Add data rows
        print('asfaf', resources_data)
        for resource in resources_data:
            role = resource.get('roleName', '')
            num_persons = str(resource.get('noOfPersons', ''))
            responsibility = resource.get('responsibility', '')
            markdown_table += f"| {role} | {num_persons} | {responsibility} |\n"
        
        print('New: ', markdown_table)
        return markdown_table
        
    def clean_html(self, html_content):
        """Clean HTML tags for markdown display"""
        if not isinstance(html_content, str):
            return html_content
            
        # Replace HTML list tags with markdown format
        cleaned_content = html_content.replace('<ul>', '')
        cleaned_content = cleaned_content.replace('</ul>', '')
        cleaned_content = cleaned_content.replace('</li>', '\n')
        cleaned_content = cleaned_content.replace('<li>', '- ')
        
        return cleaned_content
            
    def create_milestones_markdown(self, milestones_data):
        # Create markdown table header
        markdown_table = "| Milestone | Duration | Deliverables |\n"
        markdown_table += "|----------|----------|-------------|\n"
        
        # Add data rows
        for milestone in milestones_data:
            milestone_name = milestone.get('milestone', '')
            duration = milestone.get('duration', '')
            deliverables = milestone.get('deliverables', '')
            markdown_table += f"| {milestone_name} | {duration} | {deliverables} |\n"
        
        return markdown_table
        
    def create_milestones_table(self, doc, milestones_data):
        # Create table with headers
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        
        # Set header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "Milestone"
        header_cells[1].text = "Duration"
        header_cells[2].text = "Deliverables"
        
        # Add data rows
        for milestone in milestones_data:
            row_cells = table.add_row().cells
            
            # Process milestone text for bold formatting
            milestone_text = milestone.get('milestone', '')
            self.add_text_to_cell(row_cells[0], milestone_text)
            
            # Process duration text for bold formatting
            duration_text = milestone.get('duration', '')
            self.add_text_to_cell(row_cells[1], duration_text)
            
            # Process deliverables text for bold formatting
            deliverables_text = self.clean_html(milestone.get('deliverables', ''))
            self.add_text_to_cell(row_cells[2], deliverables_text)
        
        return table

    def generate_sow_document(self, sow_data, output_filename="Generated_SOW_final.docx"):
        """Generate a DOCX document from SOW data"""
        markdown = ''
        newLineChar = '\n\n'
        doc = docx.Document()
        
        # Set default font size for normal style
        style = doc.styles['Normal']
        font = style.font
        font.size = Pt(12)
        
        # Add logo at the beginning of the document
        try:
            response = requests.get(LOGO_URL)
            if response.status_code == 200:
                logo_stream = BytesIO(response.content)
                doc.add_picture(logo_stream, width=Inches(3))
            else:
                print(f"⚠️ Failed to fetch logo from URL: {LOGO_URL}. Status code: {response.status_code}")
        except Exception as e:
            print(f"⚠️ Error adding logo: {str(e)}")
        
        # Add project name as heading with custom size
        heading = doc.add_heading(sow_data["Project Name"], level=1)
        # Set font size to 16pt for all runs in the heading
        for run in heading.runs:
            run.font.size = Pt(16)
        doc.add_paragraph('')
        markdown += f'## {sow_data["Project Name"]}\n\n'
        # {sow_data['Company Information']}
        # sow_data['Client']
        intro = (f"This Statement of Work (\"SOW\") is made and entered into as of {sow_data['SOW Effective Date']} (the \"Effective Date\") by and between:\n\n"
        f"1. {sow_data['Company Name']}, a corporation organized and existing under the laws of [Jurisdiction], having its principal "
        f"place of business at [Address] (hereinafter referred to as \"Provider\"); and\n\n"
        f"2. {sow_data['Client Name'] or 'Client'}, a [business entity type] organized and existing under the laws of [Jurisdiction], having its principal "
        f"place of business at [Address] (hereinafter referred to as \"Client\").\n\n"
        f"This SOW is executed pursuant to the Master Services Agreement between Provider and Client dated {sow_data['Agreement Date']} "
        f"(the \"Agreement\"). In the event of any conflict between this SOW and the Agreement, the terms of the Agreement shall "
        f"govern unless explicitly stated otherwise in this SOW with specific reference to the Agreement provisions being modified.")
        self.add_paragraph(doc, intro)
        markdown += f"{intro}{newLineChar}"
        
        section_order = [
            "Services Description", "Deliverables", "Milestones", "Acceptance",
            "Personnel and Locations", "Representatives", "Client Representatives",
            "Contractor Resources", "Terms & Conditions", "Fees", "Expenses", "Taxes", "Conversion",
            "Limitation of Liability", "Service Level Agreement", "Assumptions", "Change Process"
        ]
        for section in section_order:
            doc.add_heading(section, level=1)
            markdown += f"## {section}{newLineChar}"
            try:
                data = sow_data[section]
                print('Section: #', section)
                if data and isinstance(data, list):
                    if section == "Milestones":
                        # Create table for milestones
                        milestones_data = sow_data.get(section, [])
                        markdown += self.create_milestones_markdown(milestones_data) + newLineChar
                        self.create_milestones_table(doc, milestones_data)
                    else:
                        # Create table for contractor resources or other list data
                        resources_data = sow_data.get(section, [])
                        markdown += self.create_contractor_resources_markdown(resources_data) + newLineChar
                        self.create_contractor_resources_table(doc, resources_data)
                else:
                    self.add_paragraph(doc, data)
                    markdown += f"{data}{newLineChar}"
            except Exception as e:
                print('Oops, we have some exception: ', e)
                self.add_paragraph(doc, '')
                markdown += newLineChar 

        # Signature section
        doc.add_heading("Signatures", level=1)
        markdown += f"## Signatures{newLineChar}"
        self.add_paragraph(doc, "IN WITNESS WHEREOF, the Parties have executed this Agreement as of the Effective Date.")
        markdown += f"IN WITNESS WHEREOF, the Parties have executed this Agreement as of the Effective Date.{newLineChar}"
        
        self.add_paragraph(doc, f"For {sow_data['Client Name'] or 'Client'} (Client)")
        markdown += f"For {sow_data['Client Name'] or 'Client'} (Client) {newLineChar}"

        self.add_paragraph(doc, f"Name: ___________________")
        self.add_paragraph(doc, f"Signature: ________________")
        markdown += f"Name: ___________________{newLineChar}"
        markdown += f"Signature: ________________{newLineChar}"
    
        self.add_paragraph(doc, f"For {sow_data['Company Name'] or ''} (Service Provider)")
        markdown += f"For {sow_data['Company Name'] or ''} (Service Provider){newLineChar}"

        self.add_paragraph(doc, f"Name: ___________________")
        self.add_paragraph(doc, f"Signature: ________________")
        markdown += f"Name: ___________________{newLineChar}"
        markdown += f"Signature: ________________{newLineChar}"

        # Save document
        static_folder = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'static')
        if not os.path.exists(static_folder):
            os.makedirs(static_folder)

        output_path = os.path.join(static_folder, output_filename)
        doc.save(output_path)

        print(f"✅ SOW document generated: {output_filename}")
        return {"fileName": output_filename, "formatted_sow_md": markdown}
        
    def process_sow(self, state):
        """Process SOW for formatting"""
        try:
            output_file = self.generate_sow_document(state['validated_sow'])
            state['doc_file_path'] = output_file['fileName']
            state['formatted_sow'] = output_file['formatted_sow_md']
            return state
        except Exception as e:
            state['error'] = f"Document formatting failed: {str(e)}"
            return state
    
    def add_text_to_cell(self, cell, text):
        """Add text to a table cell with handling for markdown bold formatting"""
        # Clear existing text
        cell.text = ""
        
        # Check if there are any markdown bold patterns (**text**)
        if "**" in text:
            paragraph = cell.paragraphs[0]
            parts = text.split("**")
            
            # Parts will alternate between normal text and bold text
            for i, part in enumerate(parts):
                if part:  # Skip empty parts
                    # Even indices are normal text, odd indices are bold text
                    run = paragraph.add_run(part)
                    if i % 2 == 1:  # This is text that was between ** and should be bold
                        run.bold = True
        else:
            # No markdown formatting, add as normal
            cell.text = text